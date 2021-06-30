from datetime import datetime
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Mm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocMaker:

    def __init__(self, config, doc_type, destination):
        self.current_row = 0
        self.document = Document()
        self.doc_type = doc_type
        self.destination = destination
        self.file_name = f'report_{datetime.now().strftime("%d-%b-%Y-%H-%M-%S")}.docx'
        self.font_name = config['font_name']
        self.header_font_size = config['header_font_size']
        self.table_header_font_size = config['table_header_font_size']
        self.table_headers = ['Affected Hosts / Rating', 'Description', 'Impact', 'Remediation']
        self.default_color = config['default_color']
        self.table_header_color = config['table_header_color']
        self.page_height = config['page_height']
        self.page_width = config['page_width']
        self.risk_colors = {
            'Critical': (233, 40, 65),
            'High': (255, 125, 30),
            'Medium': (255, 192, 0),
            'Low': (0, 176, 80)
        }

    def process(self, data):
        self.set_page_params()
        self.create_doc_header()
        self.create_table(data)
        self.save_document()

    def set_page_params(self):
        section = self.document.sections[0]
        section.page_height = Mm(self.page_height)
        section.page_width = Mm(self.page_width)

    def create_doc_header(self):
        header_1 = self.document.add_heading('Vulnerability Details and Mitigation ', 0)
        self.document.add_paragraph()
        self.document.add_paragraph().add_run().add_break()
        header_1_style = header_1.style
        header_1_style.font.name = self.font_name
        header_1_style.font.size = Pt(self.header_font_size)
        header_1_style.font.color.rgb = RGBColor(*self.default_color)
        header_1.alignment = 0

    def create_table_headers(self, name, risk):
        paragraph = self.document.add_paragraph()
        run_1 = paragraph.add_run(name)
        run_2 = paragraph.add_run(f' ({risk})')
        for run in paragraph.runs:
            run.bold = True
            run.font.name = self.font_name
            run.font.size = Pt(self.table_header_font_size)
        run_1.font.color.rgb = RGBColor(*self.table_header_color)
        run_2.font.color.rgb = RGBColor(*self.risk_colors.get(risk))

    def create_new_table(self, row=1, columns=4):
        table = self.document.add_table(rows=row, cols=columns, style="Table Grid")
        self.set_cell_styling(table.rows[self.current_row].cells, *self.table_headers, bold=True,
                              color=self.default_color)
        self.set_table_styling(table, 'top', 'left', 'right')
        self.current_row += 1
        return table

    @staticmethod
    def set_table_styling(table, *args):
        tbl = table._tbl
        cell_number = 0
        for cell in tbl.iter_tcs():
            x = ['top', 'bottom', 'left', 'right']
            tc_pr = cell.tcPr
            tc_borders = OxmlElement("w:tcBorders")
            for border in args:
                side = OxmlElement(f'w:{border}')
                side.set(qn("w:val"), "nil")
                tc_borders.append(side)
            if cell_number > 3:
                for i in set(x).difference(set(args)):
                    side = OxmlElement(f'w:{i}')
                    side.set(qn("w:val"), "single")
                    side.set(qn("w:sz"), "12")
                    side.set(qn("w:color"), "4f2d7f")
                    tc_borders.append(side)
            cell_number += 1
            tc_pr.append(tc_borders)

    def set_cell_styling(self, cells, *args, color=None, bold=None):
        for index, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                run = paragraph.add_run(args[index])
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                if color:
                    run.font.color.rgb = RGBColor(*color)
                if bold:
                    run.bold = True
                run.font.name = self.font_name
                run.font.size = Pt(self.table_header_font_size)

    def fill_in_the_table(self, table, *args):
        table.add_row()
        table.alignment = 0
        cells = table.rows[self.current_row].cells
        self.set_cell_styling(cells, *args)
        self.current_row = 0

    def create_table(self, data):
        for name in data:
            target_data = data[name]
            risk = target_data['Risk']
            cve = ', '.join(target_data['CVE']) if target_data['CVE'] else ""
            description = f'{target_data["Synopsis"]}'
            if cve:
                description += f"\n({cve})"
            remediation = target_data['Solution']
            impact = target_data['Description'].replace('\n', '')
            network = ''
            self.create_table_headers(name, risk)
            table = self.create_new_table()
            for n in target_data['Network']:
                network += f'{n}\n'
            self.fill_in_the_table(table, network, description, impact, remediation)
            self.document.add_paragraph()
            self.set_table_styling(table, 'left', 'right')

    def save_document(self):
        self.document.save(f"{self.destination}{self.file_name}")
